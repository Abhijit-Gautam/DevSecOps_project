import os
import re
import docx
import fitz  # PyMuPDF
import time
import logging
import torch
import requests
import unicodedata

#import pythoncom
import uuid
import difflib
from flask import send_from_directory

from werkzeug.utils import secure_filename
from pdf2docx import Converter

#from docx2pdf import convert
from flask import Flask, request, render_template_string, redirect, url_for
from typing import Dict, Tuple, Optional, List
from functools import lru_cache
from datetime import datetime  # Make sure this is at the top
from threading import Lock
import uuid
from gpt4all import GPT4All
from transformers import RobertaTokenizer, RobertaForSequenceClassification

import pandas as pd

EXCEL_LOG_PATH = "evaluated_reports.xlsx"


def truncate_feedback(text, max_words=65):
    words = text.strip().split()
    if len(words) <= max_words:
        truncated = ' '.join(words)
    else:
        truncated = ' '.join(words[:max_words + 10])  # allow buffer
        if '.' in truncated:
            truncated = truncated[:truncated.rfind('.') + 1]
        else:
            truncated = ' '.join(words[:max_words]) + '.'
    return truncated.strip()



# ========== Configuration ==========
app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
REPORTS_FOLDER = 'downloads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# API Keys (Should be moved to environment variables in production)

ELSEVIER_API_KEY = "01d4c17c24af98cc5f6ce9526a7087ec"


# Model Configuration
MODEL_NAME = "mistral-7b-openorca.gguf2.Q4_0.gguf"
MODEL_PATH = "models"
ROBERTA_MODEL_PATH = r"/home/cselab3/Desktop/Academic Report Evaluator/roberta-base/best_model" 
DEVICE = 'cuda' if torch.cuda.is_available() else 'cpu'
MAX_CONTENT_LENGTH = 10 * 1024 * 1024  # 10MB
CONFIDENCE_THRESHOLD = 0.8
FEEDBACK_MAX_LENGTH = 500

# Section definitions for evaluation (5 main sections)
EVALUATION_SECTIONS = {
    "abstract": [
        "abstract", "executive summary", "summary"
    ],
    "introduction": [
        "introduction", "intro", "problem statement", "objectives",
        "general introduction", "motivation", "chapter 1"
    ],
    "methodology": [
        "methodology", "development environment", "software design",
        "system design", "design", "implementation", "experimental setup",
        "research methodology", "software requirement specifications",
        "assignment work", "algorithm", "procedure", "experiment", "code"
    ],
    "results": [
        "results", "result analysis", "findings", "observations",
        "experimental results", "performance analysis", "data analysis",
        "results and discussion", "results & discussion", "output",
        "discussion", "graphs", "result interpretation"
    ],
    "conclusion": [
        "conclusion", "concluding remarks", "summary and conclusion",
        "discussion and conclusion", "results and conclusion",
        "conclusion & scope for future work", "future scope",
        "scope for future work", "concluding observations"
    ]
}


# Section definitions for semantic structure evaluation (6 sections)
SEMANTIC_SECTIONS = {
    **EVALUATION_SECTIONS,
    "literature survey": ["literature survey", "literature review", "related work", 
                         "review of literature", "background", "literature study"],
    "references": ["references", "bibliography", "works cited", "citation"]
}

# ========== Initialize Services ==========
# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log'),
        logging.StreamHandler()
    ]
)

# Initialize models
try:
    roberta_tokenizer = RobertaTokenizer.from_pretrained(ROBERTA_MODEL_PATH)
    roberta_model = RobertaForSequenceClassification.from_pretrained(ROBERTA_MODEL_PATH)
    roberta_model.to(DEVICE)
    roberta_model.eval()
    logging.info("RoBERTa model loaded successfully")
except Exception as e:
    logging.error(f"Failed to load RoBERTa model: {e}")
    raise

try:
    llm = GPT4All(model_name=MODEL_NAME, model_path=MODEL_PATH, allow_download=False)
    llm_lock = Lock()
    logging.info("LLM model loaded successfully")
except Exception as e:
    logging.error(f"Failed to initialize LLM: {e}")
    llm = None




# <<< Add the validation function here >>>
def is_valid_section(text, min_words=5):
    return isinstance(text, str) and len(text.strip().split()) >= min_words



# ========== Helper Functions ==========
def allowed_file(filename: str) -> bool:
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'pdf', 'docx'}

def convert_docx_to_pdf(docx_path: str) -> str:
    pythoncom.CoInitialize()
    pdf_path = docx_path.replace(".docx", ".converted.pdf")
    convert(docx_path, pdf_path)
    return pdf_path
def safe_ascii(val):
    if not isinstance(val, str):
        val = str(val)
    return unicodedata.normalize('NFKD', val).encode('ascii', 'ignore').decode('ascii')
def convert_pdf_to_docx(pdf_path: str) -> str:
    docx_path = pdf_path.replace(".pdf", ".docx")
    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None)
    cv.close()
    return docx_path

def extract_text_from_docx(file_path: str) -> Tuple[str, docx.Document]:
    doc = docx.Document(file_path)
    text = "\n".join([p.text for p in doc.paragraphs])
    return text, doc


def extract_text_from_pdf(file_path: str) -> Tuple[str, int]:
    text, blank_pages = [], 0
    with fitz.open(file_path) as pdf:
        for page in pdf:
            page_text = page.get_text().strip()
            if len(page_text) < 30:
                blank_pages += 1
            text.append(page_text)
    return "\n".join(text), blank_pages
def extract_title(doc):
    # Simple: Return the first non-empty paragraph as title (customize as needed)
    for para in doc.paragraphs:
        if para.text.strip():
            return para.text.strip()
    return "Untitled"

def chunk_text_for_roberta(text: str, tokenizer, chunk_size: int = 450, overlap: int = 50) -> List[str]:
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        chunk = " ".join(words[i:i+chunk_size])
        tokenized = tokenizer(chunk, truncation=True, padding=False, return_tensors="pt")
        if tokenized['input_ids'].shape[1] <= 512:
            chunks.append(chunk)
        i += chunk_size - overlap
    return chunks


def detect_page_numbers_pdf(path: str) -> bool:
    import os
    if not os.path.exists(path):
        print("File does not exist:", path)
        return False
    if not path.lower().endswith(".pdf"):
        #print("File is not a PDF:", path)
        return False
    try:
        number_only_pattern = re.compile(r'^\s*(page\s*)?(\d+|[ivxlc]+)\s*$', re.IGNORECASE)
        with fitz.open(path) as pdf:
            for page in pdf:
                blocks = page.get_text("dict")["blocks"]
                height = page.rect.height
                for block in blocks:
                    if block["type"] != 0:
                        continue
                    for line in block["lines"]:
                        y = line["bbox"][1]
                        if y < height * 0.15 or y > height * 0.85:
                            text = " ".join(span["text"] for span in line["spans"]).strip()
                            if number_only_pattern.match(text):
                                return True
        return False
    except Exception as e:
        print(f"[ERROR] Could not open PDF '{path}': {e}")
        return False

def clean_text(text: str) -> str:
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\[[^\]]+\]', '', text)
    text = re.sub(r'\([^)]+\)', '', text)
    return text.strip()
import difflib

def is_section_heading(text, aliases):
    """Return True if text matches any alias (with optional numbering, colon, etc)."""
    text = text.strip().lower()
    # Remove chapter numbers and symbols like "CHAPTER-3:", "3.", "III:"
    text_clean = re.sub(r'^(chapter|ch)?\s*[\divxlc]+\s*[:.\-\–—]*', '', text).strip()
    for alias in aliases:
        alias_clean = alias.lower()
        # Allow fuzzy/loose match (e.g., "Assignment Work" ≈ "assignment work")
        if alias_clean in text_clean or difflib.SequenceMatcher(None, alias_clean, text_clean).ratio() > 0.85:
            return True
    return False

def extract_sections(text: str, doc: Optional[docx.Document] = None) -> Dict[str, str]:
    # Initialize empty sections
    sections = {sec: "" for sec in EVALUATION_SECTIONS}
    heading_to_section = {}

    # Build alias map from section aliases
    alias_map = {}
    for sec, aliases in EVALUATION_SECTIONS.items():
        for alias in aliases:
            alias_map[alias.lower()] = sec

    def find_best_section(heading: str) -> Optional[str]:
        # Remove 'Chapter 1:', 'CH II -', etc.
        heading_clean = re.sub(r'^(chapter|ch)?\s*[\divxlc]+\s*[:.\-\–—]*', '', heading.strip().lower())

        # First: Exact or fuzzy alias match
        for alias, sec in alias_map.items():
            if alias in heading_clean or difflib.SequenceMatcher(None, alias, heading_clean).ratio() > 0.85:
                # Prevent false positives like "Assignment Work" for Methodology
                if sec == "methodology" and any(bad in heading_clean for bad in ["assignment", "work done", "task", "review"]):
                    return None
                return sec

        # Fallback heuristics (only if alias didn't match)
        if any(word in heading_clean for word in ["method", "implement", "algorithm", "architecture", "design", "procedure"]):
            if not any(bad in heading_clean for bad in ["assignment", "work done", "task", "review"]):
                return "methodology"
        if any(word in heading_clean for word in ["result", "analysis", "observation", "output"]):
            return "results"
        if "conclusion" in heading_clean or "future" in heading_clean:
            return "conclusion"
        if "intro" in heading_clean or "motivation" in heading_clean:
            return "introduction"
        if "abstract" in heading_clean or "summary" in heading_clean:
            return "abstract"

        return None

    # Extract content based on whether it's from docx or plain text
    if doc:
        current_section = None
        for para in doc.paragraphs:
            ptext = para.text.strip()
            if not ptext:
                continue
            if len(ptext) <= 100:  # Short enough to be a heading
                best_section = find_best_section(ptext)
                if best_section:
                    current_section = best_section
                    continue  # skip heading line itself
            if current_section:
                sections[current_section] += ptext + "\n"
    else:
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        current_section = None
        for line in lines:
            best_section = find_best_section(line)
            if best_section:
                current_section = best_section
                continue
            if current_section:
                sections[current_section] += line + " "

    return {k: v.strip() for k, v in sections.items()}





def extract_reference_section(text: str) -> str:
    lines = text.splitlines()
    keep, collected = False, []
    for line in lines:
        lower = line.lower().strip()
        if any(sec in lower for sec in ['references', 'bibliography']):
            keep = True
        if keep and not any(sec in lower for sec in ['appendix', 'acknowledgement']):
            collected.append(line)
    return '\n'.join(collected)




# ========== Evaluation Functions ==========
def score_section(title: str, domain: str, section_name: str, section_text: str) -> Tuple[int, float]:
    input_text = f"TITLE: {title} | DOMAIN: {domain} | SECTION: {section_name} | CONTENT: {section_text}"
    encoded = roberta_tokenizer(
        input_text,
        truncation=True,
        padding="max_length",
        max_length=512,
        return_tensors="pt"
    ).to(DEVICE)
    with torch.no_grad():
        logits = roberta_model(**encoded).logits
        prob = torch.softmax(logits, dim=-1)
        conf, pred = torch.max(prob, dim=-1)
    return int(pred.item()), float(conf.item())

@lru_cache(maxsize=100)
def predict_domain(title: str, abstract: str) -> str:
    prompt = f"""
You are a strict research domain classifier.

Your task is to output ONLY the exact technical domain (like: Deep Learning, Cybersecurity, Natural Language Processing, Cloud Computing, Operating Systems, Data Mining, IoT, etc.) based on the given project TITLE and ABSTRACT.

❗️Output must:
- Contain only the domain name.
- Be a single line with no extra words or explanation.
- Avoid phrases like “Based on...” or “The domain is:”.
- Do not include colons, punctuation, or sentences.
- Just output the domain label — e.g., Cybersecurity

Title: {title.strip()}
Abstract: {abstract.strip()[:700]}

Domain:
"""

    try:
        response = llm.generate(prompt, max_tokens=10, temp=0.1)
        print(f"[DOMAIN RAW LLM RESPONSE]: {repr(response)}")

        # Extract first non-empty line and clean it
        domain_line = next((line for line in response.strip().splitlines() if line.strip()), "")
        domain = domain_line.strip().split(":")[-1].strip()  # If colon exists, keep right part
        domain = domain.strip(' "\'.')  # Remove quotes or punctuation
        return domain or "Unknown Domain"
    except Exception as e:
        logging.error(f"Domain prediction failed: {e}")
        return "Unknown Domain"



def evaluate_section_with_llm(section: str, content: str) -> Tuple[int, str]:
    section_name = section.capitalize()
    section_content = (content or "").strip()
    truncated_content = " ".join(section_content.split()[:250])
    prompt = f"""
You are a strict academic reviewer. For the section below:
1. Give a numeric score (0 = missing or irrelevant, 1 = present but weak, 2 = excellent).
2. Provide 1-2 sentences of SPECIFIC reviewer feedback to help the author improve, even if the section is missing or weak.
If the content is missing or too short, do NOT say 'automated evaluation completed'—explicitly explain what is missing and what a good section should contain.

Format exactly as:
===
|SCORE|<score>
|FEEDBACK|<feedback>
===

Section: {section_name}
Content: {truncated_content}

Your feedback must be in 2-3 sentences, maximum 60 words, and must end with a complete sentence and a full stop. Do not leave any sentence incomplete or cut off.
"""
    try:
        with llm_lock:
            response = llm.generate(prompt, max_tokens=180, temp=0.35)
        score_match = re.search(r'\|SCORE\|([0-2])', response)
        feedback_match = re.search(r'\|FEEDBACK\|(.+?)(?:[\n=]|$)', response, re.DOTALL)
        if score_match and feedback_match:
            score = int(score_match.group(1))
            feedback = truncate_feedback(feedback_match.group(1), max_words=65)
            return score, feedback
    except Exception:
        return 1, "LLM feedback unavailable. Please review section manually."


def generate_feedback_with_score(section: str, content: str, score: int) -> str:
    section_name = section.capitalize()
    section_content = (content or "").strip()
    truncated_content = " ".join(section_content.split()[:250])
    prompt = f"""
You are an academic reviewer. Given the following section content and a score assigned by an automated evaluator ({score}/2), provide specific feedback that matches this score.

- If the score is 2, give positive but still actionable feedback.
- If the score is 1, explain what is missing and give suggestions for improvement.
- If the score is 0, explain clearly what is missing and what a good section should contain.

Section: {section_name}
Content: {truncated_content}
Score: {score}/2

Respond in 2-3 sentences (maximum 60 words). Ensure your response ends with a complete sentence and a full stop. Do not leave any sentence incomplete or cut off.
"""
    with llm_lock:
        response = llm.generate(prompt, max_tokens=180, temp=0.3)
    # Always truncate at word/sentence boundary and enforce completeness
    feedback = truncate_feedback(response, max_words=65)
    return feedback



def evaluate_literature(text: str, doc: Optional[docx.Document] = None) -> Dict:
    """
    Evaluates the Literature Survey and References sections for citation quality, diversity, and relevance.
    Scoring Rules:
        - 3 points: At least 5 references AND 3+ of the first 5 are verified in Elsevier.
        - 2 points: At least 5 references AND 1-2 of the first 5 are verified in Elsevier.
        - 1 point: At least 5 references but none verifiable, OR 3-4 references present.
        - 0 points: Fewer than 3 references.
    Returns:
        {
            "score": ...,
            "feedback": [ ... ],
            "tips": [ ... ]
        }
    """
    import re
    MAX_REFERENCES = 5
    feedback = []
    tips = []

    # ---- SECTION EXTRACTION ----
    literature_text = ""
    references_text = ""
    if doc is not None:
        current = None
        for para in doc.paragraphs:
            ptext = para.text.strip()
            ptext_lower = ptext.lower()
            # Detect section start
            if any(x in ptext_lower for x in ['literature survey', 'literature review', 'related work', 'background', 'literature study']):
                current = "literature"
                continue
            elif any(x in ptext_lower for x in ['references', 'bibliography', 'works cited', 'citation']):
                current = "references"
                continue
            elif any(x in ptext_lower for x in ["appendix", "acknowledgement"]):
                current = None
            # Collect lines
            if current == "literature":
                literature_text += ptext + "\n"
            elif current == "references":
                references_text += ptext + "\n"
    else:
        # Fallback: Use regex on text
        lit_match = re.search(r'(literature (survey|review|study)|related work|background)[\s:\-]*\n(.+?)(?=(\n\S+[\s:\-]*\n)|$)', text, re.IGNORECASE | re.DOTALL)
        if lit_match:
            literature_text = lit_match.group(3).strip()
        ref_match = re.search(r'(references|bibliography|works cited|citation)[\s:\-]*\n(.+?)(?=(\n\S+[\s:\-]*\n)|$)', text, re.IGNORECASE | re.DOTALL)
        if ref_match:
            references_text = ref_match.group(2).strip()

    # ---- REFERENCE COLLECTION ----
    combined_refs = []
    if references_text:
        combined_refs.extend([line.strip() for line in references_text.split("\n") if len(line.strip()) > 10])
    if literature_text:
        # Extract possible citations from literature text as fallback
        combined_refs.extend([line.strip() for line in literature_text.split("\n") if (re.search(r'\d{4}|et al\.|http', line) and len(line.strip()) > 10)])
    # Deduplicate and cap to first 5
    ref_lines = list(dict.fromkeys([ref for ref in combined_refs if ref]))  # remove duplicates, preserve order

    total_refs = len(ref_lines)
    refs_to_check = ref_lines[:MAX_REFERENCES]

    # ---- VERIFICATION USING ELSEVIER ----
    verified = 0
    for line in refs_to_check:
        # Try to extract the main title part (strip numbers, links, etc.)
        title = re.sub(r'http\S+|\d+\.\s*|pp\.|vol\.|no\.|[^\w\s]', '', line.split(":")[-1]).strip()
        if len(title.split()) >= 3:
            try:
                if check_elsevier_match(title):
                    verified += 1
            except Exception:
                continue

    # ---- SCORING ----
    if total_refs < 3:
        score = 0
        feedback.append("❌ Fewer than 3 references found.")
    elif total_refs < 5:
        score = 1
        feedback.append("⚠ At least 5 references are recommended for a strong literature review.")
    elif verified >= 3:
        score = 3
        feedback.append(f"✅ {verified}/5 references verified in global databases.")
    elif verified >= 1:
        score = 2
        feedback.append(f"⚠ {verified}/5 references verified in global databases.")
    else:
        score = 1
        feedback.append("❌ References found but not verifiable in academic sources.")

    # ---- TIPS LOGIC ----
    if total_refs < 5:
        tips.append("Include at least 5 academic references to strengthen your review.")
    if verified < 3 and total_refs >= 5:
        tips.append("Ensure cited works are from peer-reviewed journals or major conferences.")
    if not literature_text.strip():
        tips.append("Summarize key contributions from prior research in the Literature Survey section.")
    elif len(literature_text.split()) < 50:
        tips.append("Expand your Literature Survey to discuss methods, findings, and limitations from more sources.")
    if not any("doi" in ref.lower() or "https://doi.org" in ref.lower() for ref in ref_lines):
        tips.append("Provide DOIs or official publisher links for cited articles.")
    tips = list(dict.fromkeys(tips))[:3]

    return {"score": score, "feedback": feedback, "tips": tips}


def check_elsevier_match(title: str) -> bool:
    url = f"https://api.elsevier.com/content/search/scopus?query=TITLE(\"{title}\")"
    headers = {'X-ELS-APIKey': ELSEVIER_API_KEY}
    try:
        r = requests.get(url, headers=headers, timeout=10)
        return r.json().get('search-results', {}).get('entry', []) != []
    except Exception:
        return False


def detect_font_consistency(doc):
    """
    Returns (score, feedback)
    - Score = 1 if all fonts are standard (Times New Roman, Calibri, Arial), else 0.
    - Feedback explains result.
    """
    standard_fonts = {"times new roman", "calibri", "arial"}
    found_fonts = set()
    # 1. Check all runs in all paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font and run.font.name:
                found_fonts.add(run.font.name.lower())
    # 2. If none found (style used), fallback to 'Normal' style
    if not found_fonts:
        try:
            style_font = doc.styles['Normal'].font.name
            if style_font:
                found_fonts.add(style_font.lower())
        except Exception:
            pass
    # 3. Evaluate
    if found_fonts & standard_fonts:
        score = 1
        feedback = (
            "The document maintains a professional appearance with consistent use of standard fonts "
            "(Times New Roman, Calibri, or Arial)."
        )
    else:
        score = 0
        feedback = (
            "Inconsistent or non-standard fonts were detected. "
            "Please use a standard font such as Times New Roman, Calibri, or Arial throughout the document."
        )
    return score, feedback

def evaluate_semantics(doc: docx.Document, text: str, blank_pages: int = None, file_path: str = None) -> Dict:
    """
    Evaluates document semantic quality with reviewer-style feedback for each parameter.
    All parameters, including Section Structure, have max score 1.
    Returns:
        {
            "total_score": ...,
            "max_score": ...,
            "details": {
                "<parameter>": {"score": ..., "max": ..., "feedback": [<reviewer-feedback>]}
            }
        }
    """
    results = {}
    total_score = 0

    # 1. Section Structure (1 mark)
    detected_sections = set()
    for para in doc.paragraphs:
        if not para.text.strip() or len(para.text.split()) > 15:
            continue
        lower_text = para.text.lower()
        for sec, aliases in SEMANTIC_SECTIONS.items():
            if any(alias in lower_text for alias in aliases):
                detected_sections.add(sec)
    missing = [sec.replace("_", " ").capitalize() for sec in SEMANTIC_SECTIONS if sec not in detected_sections]
    if not missing:
        section_score = 1
        section_feedback = "All major sections are clearly present and properly titled."
    else:
        section_score = 0
        section_feedback = (
            f"The following section(s) are missing or unclear: {', '.join(missing)}. "
            f"Please ensure all standard report sections are included and properly labeled."
        )
    results["section_structure"] = {"score": section_score, "max": 1, "feedback": [section_feedback]}
    total_score += section_score

    # 2. Font Consistency (1 mark)
    font_score, font_feedback = detect_font_consistency(doc)
    results["font_consistency"] = {"score": font_score, "max": 1, "feedback": [font_feedback]}
    total_score += font_score

    # 3. Line Spacing (1 mark) -- ROBUST LOGIC
    spacings = []
    for p in doc.paragraphs:
        spacing = p.paragraph_format.line_spacing
        if spacing is None:
            try:
                spacing = p.style.paragraph_format.line_spacing
            except Exception:
                spacing = None
        try:
            if spacing and 1.0 <= float(spacing) <= 3.0:
                spacings.append(round(float(spacing), 1))
        except Exception:
            continue
    if spacings:
        avg_spacing = sum(spacings) / len(spacings)
        if 1.1 <= avg_spacing <= 2.1:
            spacing_score = 1
            spacing_feedback = (
                f"Line spacing appears appropriate ({round(avg_spacing,1)}x), making the document easy to read."
            )
        else:
            spacing_score = 0
            spacing_feedback = (
                f"Line spacing is inconsistent or falls outside the recommended range (detected average: {round(avg_spacing,1)}x). "
                "Please set line spacing to around 1.5 for improved readability."
            )
    else:
        spacing_score = 0
        spacing_feedback = (
            "Line spacing could not be reliably detected. "
            "Please ensure consistent 1.5x line spacing is used throughout the document."
        )
    results["line_spacing"] = {"score": spacing_score, "max": 1, "feedback": [spacing_feedback]}
    total_score += spacing_score

    # 4. Figures & Tables (1 mark)
    found_figure = bool(re.search(r'\b(figure|fig\.?)\s*\d+', text.lower()))
    found_table = bool(re.search(r'\b(table|tab\.?)\s*\d+', text.lower()))
    if found_figure and found_table:
        fig_table_score = 1
        fig_table_feedback = (
            "Both figures and tables are included and properly labeled, which strengthens the clarity of your results."
        )
    elif found_figure or found_table:
        fig_table_score = 1
        fig_table_feedback = (
            "Either figures or tables are present, but not both. Consider including and labeling both to enhance your analysis."
        )
    else:
        fig_table_score = 0
        fig_table_feedback = (
            "No labeled figures or tables were found in the document. "
            "Consider including relevant figures and tables with proper captions to support your findings."
        )
    results["figures_tables"] = {"score": fig_table_score, "max": 1, "feedback": [fig_table_feedback]}
    total_score += fig_table_score

    # 5. Blank Pages (1 mark)
    if blank_pages is None:
        paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        blank_pages = 0
        for i in range(0, len(paragraphs), 50):
            chunk = paragraphs[i:i + 50]
            words = sum(len(p.text.strip().split()) for p in chunk)
            if words < 20:
                blank_pages += 1
    if not blank_pages or blank_pages <= 1:
        blank_page_score = 1
        blank_feedback = "No unnecessary blank or nearly blank pages were detected in the document."
    else:
        blank_page_score = 0
        blank_feedback = (
            f"{blank_pages} page(s) in the document appear to be blank or nearly blank. "
            "Please review and remove any unnecessary blank pages for a cleaner submission."
        )
    results["blank_pages"] = {"score": blank_page_score, "max": 1, "feedback": [blank_feedback]}
    total_score += blank_page_score

    # 6. Citations (1 mark)
    citation_count = len(re.findall(r'\b[A-Z][a-z]+ et al\.?,? \d{4}\b|\[\d+\]', text))
    if citation_count >= 3:
        citation_score = 1
        citation_feedback = (
            "Adequate in-text citations are present, supporting the discussion with external sources."
        )
    else:
        citation_score = 0
        citation_feedback = (
            "Few or no in-text citations were found in the document. "
            "Be sure to reference all important works and cite sources in the main text."
        )
    results["citations"] = {"score": citation_score, "max": 1, "feedback": [citation_feedback]}
    total_score += citation_score

    # 7. Page Numbers (1 mark)
    page_num_score = 1 if (file_path and detect_page_numbers_pdf(file_path)) else 0
    if page_num_score:
        page_num_feedback = "Page numbers are clearly visible in headers or footers, making the report easy to navigate."
    else:
        page_num_feedback = (
            "No page numbers were detected in the document. "
            "Please include page numbers in the header or footer for better organization."
        )
    results["page_numbers"] = {"score": page_num_score, "max": 1, "feedback": [page_num_feedback]}
    total_score += page_num_score

    return {
        "total_score": min(total_score, 7),
        "max_score": 7,
        "details": results
    }






def generate_report_file(results, out_dir=REPORTS_FOLDER):
    os.makedirs(out_dir, exist_ok=True)
    filename = f"report_{uuid.uuid4().hex[:8]}.txt"
    file_path = os.path.join(out_dir, filename)
    with open(file_path, "w", encoding="utf-8") as f:
        f.write("Academic Report Evaluation Summary\n")
        f.write(f"Domain: {results['domain']}\n")
        f.write(f"Total Score: {results['total_score']}/{results['max_score']}\n\n")
        for sec, score in results["section_scores"].items():
            f.write(f"{sec.capitalize()}: {score}/2\nFeedback: {results['section_feedback'][sec]}\n\n")
        f.write(f"Literature Review: {results['literature']['score']}/3\n")
        for fb in results['literature']['feedback']:
            f.write(f"- {fb}\n")
        f.write(f"\nDocument Quality: {results['semantics']['total_score']}/7\n")
        for k, v in results['semantics']['details'].items():
            f.write(f"{k.replace('_',' ').capitalize()}: {v['score']}/{v['max']}\n")
            for fb in v['feedback']:
                f.write(f"  - {fb}\n")
    
    return filename

def update_excel_log(filename: str, metadata: Dict):
    """
    Appends evaluation details to an Excel log.
    Creates the file if it doesn't exist.
    """
    row = {
        "Filename": filename,
        "Domain": metadata.get("domain"),
        "Abstract Score": metadata.get("abstract_score"),
        "Introduction Score": metadata.get("introduction_score"),
        "Methodology Score": metadata.get("methodology_score"),
        "Results Score": metadata.get("results_score"),
        "Conclusion Score": metadata.get("conclusion_score"),
        "Total Section Score": metadata.get("total_section_score"),
        "Literature Score": metadata.get("literature_score"),
        "Semantic Score": metadata.get("semantic_score"),

        "Evaluation Date": metadata.get("evaluation_date")
    }

    try:
        # Load existing Excel file or create new
        if os.path.exists(EXCEL_LOG_PATH):
            df_existing = pd.read_excel(EXCEL_LOG_PATH)
            df_updated = pd.concat([df_existing, pd.DataFrame([row])], ignore_index=True)
        else:
            df_updated = pd.DataFrame([row])

        df_updated.to_excel(EXCEL_LOG_PATH, index=False)
        print(f"📄 Excel log updated: {EXCEL_LOG_PATH}")
    except Exception as e:
        print(f"❌ Failed to update Excel log: {e}")


# ========== Flask Routes ==========
@app.route('/')
def home():
    return render_template_string(HOME_TEMPLATE)

# ... (other code above remains the same)

@app.route('/evaluate', methods=['POST'])
def evaluate():
    debug_trace = []
    debug_trace.append("Starting evaluation process...")
    logging.info("Starting evaluation process...")

    if 'report' not in request.files:
        debug_trace.append("Error: No file uploaded")
   
        return render_template_string(HOME_TEMPLATE, error="No file uploaded")

    file = request.files['report']
    if not file.filename:
        debug_trace.append("Error: Empty file")
 
        return render_template_string(HOME_TEMPLATE, error="Empty file")

    filename = secure_filename(file.filename)
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    file.save(file_path)
    debug_trace.append(f"File saved: {filename}")
    logging.info(f"File saved: {filename}")

    temp_files = [file_path]
    text = ""
    doc = None
    blank_pages = None
    page_number_detected = False
    pdf_path = None
    docx_path = None

    try:
        if filename.lower().endswith('.docx'):
            debug_trace.append("Processing DOCX file...")
         
            text, doc = extract_text_from_docx(file_path)
            title = extract_title(doc)
            blank_pages = None
            debug_trace.append(f"Extracted text length: {len(text)} chars")
            debug_trace.append(f"First 100 chars: {text[:100]!r}")

            try:
                pdf_path = convert_docx_to_pdf(file_path)
                temp_files.append(pdf_path)
                if not os.path.exists(pdf_path) or os.path.getsize(pdf_path) == 0:
                    debug_trace.append("DOCX to PDF conversion failed.")
                  
                    return render_template_string(HOME_TEMPLATE, error="DOCX to PDF conversion failed.")
                try:
                    page_number_detected = detect_page_numbers_pdf(pdf_path)
                except Exception as e:
                    page_number_detected = False
                    debug_trace.append(f"PDF page number check failed: {e}")
                    
            except Exception as e:
                debug_trace.append(f"Error in DOCX to PDF conversion: {e}")
                
                page_number_detected = False

        elif filename.lower().endswith('.pdf'):
            debug_trace.append("Processing PDF file...")
           
            try:
                text, blank_pages = extract_text_from_pdf(file_path)
                title = os.path.splitext(filename)[0]
                debug_trace.append(f"Extracted text length: {len(text)} chars, blank pages: {blank_pages}")
                debug_trace.append(f"First 100 chars: {text[:100]!r}")
            except Exception as e:
                debug_trace.append(f"PDF extraction error: {e}")
              
                return render_template_string(HOME_TEMPLATE, error="Could not extract text from PDF.")

            try:
                docx_path = convert_pdf_to_docx(file_path)
                temp_files.append(docx_path)
                if not os.path.exists(docx_path) or os.path.getsize(docx_path) == 0:
                    debug_trace.append("PDF to DOCX conversion failed.")
                   
                    return render_template_string(HOME_TEMPLATE, error="PDF to DOCX conversion failed.")
                _, doc = extract_text_from_docx(docx_path)
            except Exception as e:
                debug_trace.append(f"PDF to DOCX conversion failed: {e}")
                
                return render_template_string(HOME_TEMPLATE, error="PDF to DOCX conversion failed.")

            try:
                page_number_detected = detect_page_numbers_pdf(file_path)
            except Exception as e:
                page_number_detected = False
                debug_trace.append(f"PDF page number check failed: {e}")
               

        else:
            debug_trace.append("Error: Unsupported file type")
            
            return render_template_string(HOME_TEMPLATE, error="Unsupported file type")

        if not text.strip():
            debug_trace.append("Error: Empty document")
            
            return render_template_string(HOME_TEMPLATE, error="Empty document")

        lines = [line.strip() for line in text.splitlines() if len(line.strip()) > 10]
        debug_trace.append(f"Extracted title: {title}")
        

        sections = extract_sections(text, doc)
        for sec, content in sections.items():
            clean_preview = content[:120].replace('\n', ' ')
            debug_trace.append(f"EXTRACTED SECTION: {sec.upper()} -> {clean_preview}...")
            

        if not is_valid_section(title, 3):
            
            return render_template_string(HOME_TEMPLATE, error="Title missing or too short.")
        if not is_valid_section(sections.get("abstract", ""), 5):
           
            return render_template_string(HOME_TEMPLATE, error="Abstract missing or too short.")

        debug_trace.append(f"DOMAIN PREDICTION INPUT: Title='{title[:80]}' Abstract='{sections.get('abstract','')[:80]}'")
        domain = predict_domain(title, sections.get("abstract", ""))
        debug_trace.append(f"Predicted domain: {domain}")
        

        section_scores = {}
        section_feedback = {}
        section_approach = {}

        for section in EVALUATION_SECTIONS:
            content = sections.get(section, "")
            approach = "Hybrid (RoBERTa)"
            if not content.strip():
                section_scores[section] = 0
                section_feedback[section] = f"{section.capitalize()} section missing or not detected."
                section_approach[section] = "Not scored"
                debug_trace.append(f"{section.upper()} | No content detected. Assigned score=0, Not scored.")
                logging.info(f"{section.upper()} | No content detected. Assigned score=0, Not scored.")
                continue

            chunks = chunk_text_for_roberta(content, roberta_tokenizer)
            chunk_scores = []
            chunk_confs = []

            for chunk in chunks:
                try:
                    s, c = score_section(title, domain, section, chunk)
                    chunk_scores.append(s)
                    chunk_confs.append(c)
                except Exception as e:
                    debug_trace.append(f"{section.upper()} | Exception in score_section: {e}")
                    logging.info(f"{section.upper()} | Exception in score_section: {e}")

            if chunk_scores:
                avg_score = round(sum(chunk_scores) / len(chunk_scores))
                avg_conf = sum(chunk_confs) / len(chunk_confs)
                score, conf = avg_score, avg_conf
            else:
                score, conf = 1, 0.0

            debug_trace.append(f"{section.upper()} | RoBERTa score={score}, conf={conf:.2f}")
            logging.info(f"{section.upper()} | RoBERTa score={score}, conf={conf:.2f}")
            if conf < CONFIDENCE_THRESHOLD:

                debug_trace.append(f"{section.upper()} | Fallback to Mistral for scoring")
                logging.info(f"{section.upper()} | Fallback to Mistral for scoring")
                try:
                    llm_score, _ = evaluate_section_with_llm(section, content)
                    llm_feedback = generate_feedback_with_score(section, content, llm_score)
                except Exception as ex:
                    llm_score, llm_feedback = 1, "LLM feedback unavailable (exception)."
                    debug_trace.append(f"{section.upper()} | Exception in LLM scoring: {ex}")
                    logging.info(f"{section.upper()} | Exception in LLM scoring: {ex}")

                section_scores[section] = llm_score
                section_feedback[section] = llm_feedback
                section_approach[section] = "LLM only"
                logging.info(f"{section.upper()} | Used Mistral for both score and feedback (fallback).")
            else:
                try:
                    llm_feedback = generate_feedback_with_score(section, content, score)
                except Exception as ex:
                    llm_feedback = "LLM feedback unavailable (exception)."
                    debug_trace.append(f"{section.upper()} | Exception in generate_feedback_with_score: {ex}")
                    logging.info(f"{section.upper()} | Exception in generate_feedback_with_score: {ex}")

                section_scores[section] = score
                section_feedback[section] = llm_feedback
                section_approach[section] = "Hybrid (RoBERTa)"
                logging.info(f"{section.upper()} | Used RoBERTa for scoring, Mistral for feedback.")

            debug_trace.append(f"{section.upper()} | Feedback: {section_feedback[section][:80].replace(chr(10),' ')}")

        total_section_score = sum(section_scores.values())
        debug_trace.append(f"Total section score: {total_section_score}/10")
    

        literature = evaluate_literature(text)
        debug_trace.append(f"Literature score: {literature['score']}/3, feedback: {literature['feedback']}")
        

        semantics = evaluate_semantics(doc, text, blank_pages, file_path)
        semantics["details"]["page_numbers"] = {
            "score": 1 if page_number_detected else 0,
            "max": 1,
            "feedback": ["✅ Page numbers detected"] if page_number_detected else ["⚠ No valid page numbers detected"]
        }
        semantics["total_score"] = sum(val.get("score", 0) for val in semantics["details"].values())
        semantics["max_score"] = sum(val.get("max", 1) for val in semantics["details"].values())
        debug_trace.append(f"Semantic total score: {semantics['total_score']}/7")
        

        # ---- Plagiarism check ----
        
        plagiarism = "Not Checked"
            

              

        total_score = total_section_score + literature["score"] + semantics["total_score"]
        max_score = 10 + 3 + 7

        results = {
            "domain": domain,
            "section_scores": section_scores,
            "section_feedback": section_feedback,
            "section_approach": section_approach,
            "literature": literature,
            "semantics": semantics,
       
            "total_score": total_score,
            "max_score": max_score,
            "debug_trace": debug_trace
        }

        metadata = {
            "domain": domain,
            "abstract_score": str(section_scores.get("abstract", 0)),
            "introduction_score": str(section_scores.get("introduction", 0)),
            "methodology_score": str(section_scores.get("methodology", 0)),
            "results_score": str(section_scores.get("results", 0)),
            "conclusion_score": str(section_scores.get("conclusion", 0)),
            "total_section_score": str(total_section_score),
            "literature_score": str(literature["score"]),
            "semantic_score": str(semantics["total_score"]),
        
            "evaluation_date": datetime.now().isoformat()
        }
        update_excel_log(filename, metadata)
        for section, feedback in section_feedback.items():
            metadata[f"{section}_feedback"] = safe_ascii(feedback[:200])

        

        

        report_filename = generate_report_file(results)
        return render_template_string(HOME_TEMPLATE, results=results, now=datetime.now(), report_filename=report_filename)

    except Exception as e:
        debug_trace.append(f"Error during evaluation: {e}")
        
        return render_template_string(HOME_TEMPLATE, error=str(e))

    finally:
        for path in temp_files:
            if path and os.path.exists(path):
                try:
                    os.remove(path)
                except Exception:
                    pass




@app.route('/download/<filename>')
def download_report(filename):
    return send_from_directory(REPORTS_FOLDER, filename, as_attachment=True)


# ========== HTML Template ==========
HOME_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Academic Report Evaluator</title>
    <link href="https://cdn.jsdelivr.net/npm/tailwindcss@2.2.19/dist/tailwind.min.css" rel="stylesheet">
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0/css/all.min.css">
    <style>
        :root {
            --primary: #2563eb;
            --primary-dark: #1d4ed8;
            --success: #16a34a;
            --warning: #d97706;
            --danger: #dc2626;
            --gray-light: #f3f4f6;
            --gray-dark: #6b7280;
        }
        
        body {
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
            line-height: 1.6;
            color: #1f2937;
        }
        
        .score-badge {
            font-size: 0.75rem;
            font-weight: 600;
            padding: 0.25rem 0.5rem;
            border-radius: 0.25rem;
        }
        
        .score-0 { background-color: #fee2e2; color: #b91c1c; }
        .score-1 { background-color: #fef3c7; color: #b45309; }
        .score-2 { background-color: #dcfce7; color: #15803d; }
        .score-3 { background-color: #d1fae5; color: #065f46; }
        
        .loading-spinner { 
            animation: spin 1s linear infinite;
            display: none;
        }
        
        @keyframes spin { 
            0% { transform: rotate(0deg); } 
            100% { transform: rotate(360deg); } 
        }
        
        .file-upload {
            transition: all 0.2s ease;
        }
        
        .file-upload:hover {
            background-color: #f8fafc;
            border-color: var(--primary);
        }
        
        .file-upload.dragover {
            background-color: #eff6ff;
            border-color: var(--primary);
            box-shadow: 0 0 0 3px rgba(37, 99, 235, 0.1);
        }
        
        .progress-bar {
            height: 6px;
            background-color: #e5e7eb;
            border-radius: 3px;
            overflow: hidden;
        }
        
        .progress-fill {
            height: 100%;
            background-color: var(--primary);
            transition: width 0.3s ease;
        }
        
        .section-card {
            transition: transform 0.2s ease, box-shadow 0.2s ease;
        }
        
        .section-card:hover {
            transform: translateY(-2px);
            box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.1);
        }
        
        .tooltip {
            position: relative;
        }
        
        .tooltip-text {
            visibility: hidden;
            width: 200px;
            background-color: #374151;
            color: #fff;
            text-align: center;
            border-radius: 6px;
            padding: 5px;
            position: absolute;
            z-index: 1;
            bottom: 125%;
            left: 50%;
            transform: translateX(-50%);
            opacity: 0;
            transition: opacity 0.3s;
            font-size: 0.75rem;
            font-weight: normal;
        }
        
        .tooltip:hover .tooltip-text {
            visibility: visible;
            opacity: 1;
        }
        
        .fade-in {
            animation: fadeIn 0.5s ease-in;
        }
        
        @keyframes fadeIn {
            from { opacity: 0; }
            to { opacity: 1; }
        }
    </style>
</head>
<body class="bg-gray-50 min-h-screen">
    <div class="container mx-auto px-4 py-8 max-w-6xl">
        <!-- Header Section -->
        <header class="mb-8 text-center">
            <h1 class="text-4xl font-bold text-gray-800 mb-2">Academic Report Evaluator</h1>
            <p class="text-lg text-gray-600 max-w-2xl mx-auto">
                Upload your academic report for comprehensive evaluation and receive detailed feedback on content, structure, and academic quality.
            </p>
        </header>
        
        <!-- Upload Card -->
        <div class="bg-white rounded-xl shadow-sm border border-gray-200 p-6 mb-8">
            <div class="flex flex-col md:flex-row gap-8">
                <div class="flex-1">
                    <h2 class="text-xl font-semibold text-gray-800 mb-4">Upload Your Report</h2>
                    <form id="evaluationForm" method="post" enctype="multipart/form-data" action="/evaluate" class="space-y-4">
                        <div class="file-upload-container">
                            <label for="fileInput" class="file-upload flex flex-col items-center justify-center w-full h-40 border-2 border-dashed border-gray-300 rounded-lg cursor-pointer bg-gray-50 hover:border-blue-500 transition">
                                <div class="flex flex-col items-center justify-center pt-5 pb-6">
                                    <i class="fas fa-cloud-upload-alt text-4xl text-gray-400 mb-3"></i>
                                    <p class="mb-2 text-sm text-gray-500 font-medium">Drag & drop your file here</p>
                                    <p class="text-xs text-gray-500">or click to browse (DOCX only, max 10MB)</p>
                                </div>
                                <input id="fileInput" name="report" type="file" class="hidden" accept=".pdf,.docx" required />
                            </label>
                            <div id="fileNameDisplay" class="mt-2 text-sm text-gray-600 hidden"></div>
                        </div>
                        <button type="submit" class="w-full bg-blue-600 hover:bg-blue-700 text-white font-medium py-3 px-4 rounded-lg transition flex items-center justify-center shadow-sm">
                            <span id="submitText">Evaluate Report</span>
                            <i id="spinner" class="fas fa-spinner loading-spinner ml-2"></i>
                        </button>
                    </form>
                </div>
                <div class="flex-1 border-l border-gray-200 pl-8">
                    <h3 class="text-lg font-medium text-gray-700 mb-3">What We Evaluate</h3>
                    <ul class="space-y-3 text-sm text-gray-600">
                        <li class="flex items-start">
                            <i class="fas fa-check-circle text-green-500 mt-1 mr-2"></i>
                            <span>Content quality of key sections (Abstract, Introduction, etc.)</span>
                        </li>
                        <li class="flex items-start">
                            <i class="fas fa-check-circle text-green-500 mt-1 mr-2"></i>
                            <span>Literature review and reference quality</span>
                        </li>
                        <li class="flex items-start">
                            <i class="fas fa-check-circle text-green-500 mt-1 mr-2"></i>
                            <span>Document structure and formatting</span>
                        </li>
                        
                        <li class="flex items-start">
                            <i class="fas fa-check-circle text-green-500 mt-1 mr-2"></i>
                            <span>Overall academic quality score</span>
                        </li>
                    </ul>
                </div>
            </div>
        </div>
        
        {% if error %}
        <div class="bg-red-50 border-l-4 border-red-500 text-red-700 p-4 mb-6 rounded flex items-start">
            <i class="fas fa-exclamation-circle mt-1 mr-3 text-red-500"></i>
            <div>
                <p class="font-medium">Error Processing Request</p>
                <p class="text-sm">{{ error }}</p>
            </div>
        </div>
        {% endif %}
        
        {% if results %}
        <div class="fade-in">
            <!-- Results Header -->
            <div class="flex justify-between items-center mb-6">
                <h2 class="text-2xl font-bold text-gray-800">Evaluation Results</h2>
                <div class="flex items-center">
                    <span class="text-sm text-gray-500 mr-2">Generated:</span>
                    <span class="text-sm font-medium">{{ now.strftime('%Y-%m-%d %H:%M') }}</span>
                </div>
            </div>
            
            <!-- Overview Card -->
            <div class="bg-white rounded-xl shadow-sm border border-gray-200 p-6 mb-6">
                <h3 class="text-xl font-semibold text-gray-800 mb-4">Document Overview</h3>
                <div class="grid grid-cols-1 md:grid-cols-3 gap-4">
                    <div class="border-r border-gray-200 pr-4">
                        <p class="text-sm text-gray-500 mb-1">Research Domain</p>
                        <p class="font-medium text-lg">{{ results.domain }}</p>
                    </div>
                   
                    <div>
                        <p class="text-sm text-gray-500 mb-1">Total Score</p>
                        <div class="flex items-center">
                            <span class="font-bold text-2xl mr-2">{{ results.total_score }}/{{ results.max_score }}</span>
                            <div class="progress-bar flex-1">
                                <div class="progress-fill" style="width: {{ (results.total_score/results.max_score)*100 }}%"></div>
                            </div>
                        </div>
                    </div>
                </div>
            </div>
            
            <!-- Section Evaluation -->
            <div class="bg-white rounded-xl shadow-sm border border-gray-200 p-6 mb-6">
                <div class="flex justify-between items-center mb-4">
                    <h3 class="text-xl font-semibold text-gray-800">Section Evaluation</h3>
                    <div class="text-sm">
                        <span class="text-gray-500">Total: </span>
                        <span class="font-bold">{{ results.section_scores.values()|sum }}/10</span>
                    </div>
                </div>
                
                <div class="grid grid-cols-1 md:grid-cols-2 gap-4 mb-6">
                    {% for section, score in results.section_scores.items() %}
                    <div class="section-card bg-white border border-gray-200 rounded-lg p-4 hover:shadow-md">
                        <div class="flex justify-between items-start mb-2">
                            <h4 class="font-medium text-gray-800 capitalize">{{ section }}</h4>
                            <div class="flex items-center">
                                <span class="font-bold mr-2">{{ score }}/2</span>
                                <span class="score-badge score-{{ score }}">{{ ["Poor", "Fair", "Good"][score] }}</span>
                            </div>
                        </div>
                        <p class="text-sm text-gray-600 mb-2">{{ results.section_feedback[section] }}</p>
                        {#
                        <div class="flex justify-between items-center text-xs text-gray-500">
                            <span>Evaluation Method: {{ results.section_approach[section] }}</span>
                            <span class="tooltip">
                                <i class="fas fa-info-circle"></i>
                                <span class="tooltip-text">This section was evaluated using {{ results.section_approach[section] }} approach</span>
                            </span>
                        </div>
                        #}
                    </div>
                    {% endfor %}
                </div>
                
                <div class="bg-gray-50 p-3 rounded-lg border border-gray-200">
                    <h4 class="font-medium text-gray-700 mb-2">Section Scoring Guide</h4>
                    <div class="grid grid-cols-3 gap-2 text-xs">
                        <div class="flex items-center">
                            <span class="score-badge score-0 mr-2">0</span>
                            <span>Missing or inadequate</span>
                        </div>
                        <div class="flex items-center">
                            <span class="score-badge score-1 mr-2">1</span>
                            <span>Present but needs improvement</span>
                        </div>
                        <div class="flex items-center">
                            <span class="score-badge score-2 mr-2">2</span>
                            <span>Well-written and complete</span>
                        </div>
                    </div>
                </div>
            </div>
            
            <!-- Literature Review -->
            <div class="bg-white rounded-xl shadow-sm border border-gray-200 p-6 mb-6">
                <div class="flex justify-between items-center mb-4">
                    <h3 class="text-xl font-semibold text-gray-800">Literature Review</h3>
                    <div class="flex items-center">
                        <span class="font-bold mr-2">{{ results.literature.score }}/3</span>
                        <span class="score-badge {% if results.literature.score >= 2 %}score-3{% elif results.literature.score >= 1 %}score-1{% else %}score-0{% endif %}">
                            {{ ["Inadequate", "Basic", "Good", "Excellent"][results.literature.score] }}
                        </span>
                    </div>
                </div>
                
                <div class="space-y-3 mb-4">
                    {% for item in results.literature.feedback %}
                    <div class="flex items-start">
                        {% if "❌" in item %}
                            <i class="fas fa-times-circle text-red-500 mt-1 mr-2"></i>
                        {% elif "⚠" in item %}
                            <i class="fas fa-exclamation-triangle text-yellow-500 mt-1 mr-2"></i>
                        {% else %}
                            <i class="fas fa-check-circle text-green-500 mt-1 mr-2"></i>
                        {% endif %}
                        <p class="text-sm flex-1">{{ item.replace("❌", "").replace("⚠", "").replace("✅", "") }}</p>
                    </div>
                    {% endfor %}
                </div>
                
                {% if results.literature.tips %}
                <div class="bg-blue-50 border border-blue-200 rounded-lg p-4">
                    <h4 class="font-medium text-blue-800 mb-2 flex items-center">
                        <i class="fas fa-lightbulb mr-2"></i> Improvement Tips
                    </h4>
                    <ul class="list-disc list-inside text-sm text-blue-700 space-y-1 pl-2">
                        {% for tip in results.literature.tips %}
                        <li>{{ tip }}</li>
                        {% endfor %}
                    </ul>
                </div>
                {% endif %}
            </div>
            
            <!-- Document Quality -->
            <div class="bg-white rounded-xl shadow-sm border border-gray-200 p-6 mb-6">
                <div class="flex justify-between items-center mb-4">
                    <h3 class="text-xl font-semibold text-gray-800">Document Quality</h3>
                    <div class="flex items-center">
                        <span class="font-bold mr-2">{{ results.semantics.total_score }}/7</span>
                        <span class="score-badge {% if results.semantics.total_score >= 6 %}score-3{% elif results.semantics.total_score >= 4 %}score-1{% else %}score-0{% endif %}">
                            {{ ["Poor", "Fair", "Good", "Excellent"][results.semantics.total_score//2] }}
                        </span>
                    </div>
                </div>
                
                <div class="grid grid-cols-1 md:grid-cols-2 gap-4">
                    {% for category, data in results.semantics.details.items() %}
                    <div class="border rounded-lg p-3 {% if data.score == data.max %}border-green-200 bg-green-50{% elif data.score == 0 %}border-red-200 bg-red-50{% else %}border-yellow-200 bg-yellow-50{% endif %}">
                        <div class="flex justify-between items-center mb-1">
                            <h4 class="font-medium capitalize">{{ category.replace('_', ' ') }}</h4>
                            <span class="text-sm font-bold">{{ data.score }}/{{ data.max }}</span>
                        </div>
                        {% for fb in data.feedback %}
                        <p class="text-xs {% if data.score == data.max %}text-green-700{% elif data.score == 0 %}text-red-700{% else %}text-yellow-700{% endif %}">
                            {{ fb.replace("✅", "").replace("⚠", "") }}
                        </p>
                        {% endfor %}
                    </div>
                    {% endfor %}
                </div>
            </div>
            
            <!-- Final Summary -->
            <div class="bg-white rounded-xl shadow-sm border border-gray-200 p-6">
                <h3 class="text-xl font-semibold text-gray-800 mb-4">Final Evaluation Summary</h3>
                
                <div class="grid grid-cols-1 md:grid-cols-3 gap-6 mb-6">
                    <!-- Content Quality -->
                    <div class="border border-gray-200 rounded-lg p-4">
                        <div class="flex items-center mb-3">
                            <div class="bg-blue-100 p-2 rounded-full mr-3">
                                <i class="fas fa-file-alt text-blue-600"></i>
                            </div>
                            <h4 class="font-medium">Content Quality</h4>
                        </div>
                        <div class="text-center py-2">
                            <span class="text-3xl font-bold">{{ results.section_scores.values()|sum }}/10</span>
                            <p class="text-sm text-gray-500 mt-1">Section Evaluation</p>
                        </div>
                        <div class="mt-3">
                            <p class="text-sm text-gray-600">
                                {% if results.section_scores.values()|sum >= 8 %}
                                Your report content is well-structured and comprehensive.
                                {% elif results.section_scores.values()|sum >= 5 %}
                                Your report content is adequate but could be improved.
                                {% else %}
                                Your report content needs significant improvement.
                                {% endif %}
                            </p>
                        </div>
                    </div>
                    
                    <!-- Academic Quality -->
                    <div class="border border-gray-200 rounded-lg p-4">
                        <div class="flex items-center mb-3">
                            <div class="bg-purple-100 p-2 rounded-full mr-3">
                                <i class="fas fa-graduation-cap text-purple-600"></i>
                            </div>
                            <h4 class="font-medium">Academic Quality</h4>
                        </div>
                        <div class="text-center py-2">
                            <span class="text-3xl font-bold">{{ results.literature.score + results.semantics.total_score }}/10</span>
                            <p class="text-sm text-gray-500 mt-1">Literature + Document</p>
                        </div>
                        <div class="mt-3">
                            <p class="text-sm text-gray-600">
                                {% if results.literature.score + results.semantics.total_score >= 8 %}
                                Your report meets high academic standards.
                                {% elif results.literature.score + results.semantics.total_score >= 5 %}
                                Your report meets basic academic requirements.
                                {% else %}
                                Your report needs significant academic improvements.
                                {% endif %}
                            </p>
                        </div>
                    </div>
                    
                    <!-- Overall Score -->
                    <div class="border border-blue-200 rounded-lg p-4 bg-blue-50">
                        <div class="flex items-center mb-3">
                            <div class="bg-blue-200 p-2 rounded-full mr-3">
                                <i class="fas fa-star text-blue-700"></i>
                            </div>
                            <h4 class="font-medium text-blue-800">Overall Score</h4>
                        </div>
                        <div class="text-center py-2">
                            <span class="text-4xl font-bold text-blue-700">{{ results.total_score }}/20</span>
                            <p class="text-sm text-blue-600 mt-1">Combined Evaluation</p>
                        </div>
                        <div class="mt-3">
                            <p class="text-sm text-blue-700">
                                {% if results.total_score >= 16 %}
                                Excellent report! Minor improvements could make it outstanding.
                                {% elif results.total_score >= 12 %}
                                Good report. Address the feedback to improve further.
                                {% elif results.total_score >= 8 %}
                                Fair report. Significant improvements needed.
                                {% else %}
                                Poor report. Major revisions required.
                                {% endif %}
                            </p>
                        </div>
                    </div>
                </div>
                
                <!-- Action Buttons -->
                <div class="flex flex-col sm:flex-row gap-3 mt-6">
                   <a href="/download/{{ report_filename }}" class="flex-1 bg-blue-600 hover:bg-blue-700 text-white font-medium py-2 px-4 rounded-lg transition flex items-center justify-center" download>
                           <i class="fas fa-download mr-2"></i> Download Full Report
                   </a>

                    <button type="button"
                              class="flex-1 bg-white border border-gray-300 hover:bg-gray-50 text-gray-700 font-medium py-2 px-4 rounded-lg transition flex items-center justify-center"
                              onclick="resetEvaluationForm()">
                              <i class="fas fa-redo mr-2"></i> Evaluate Another Report
                    </button>
                    <script>
        function resetEvaluationForm() {
                             window.location.href = '/'; // Reloads the page, clears form/results
                     }
                    </script>

            
            <!-- Debug Information -->
            <div class="mt-8">
                <button onclick="toggleDebug()" class="text-gray-500 hover:text-gray-700 text-sm font-medium flex items-center">
                    <i class="fas fa-bug mr-2"></i> 
                    <span id="debugToggleText">Show Debug Information</span>
                </button>
                <div id="debugInfo" class="hidden mt-4 bg-gray-100 p-4 rounded-lg text-xs font-mono overflow-auto max-h-60 border border-gray-200">
                    <div class="font-bold text-gray-700 mb-2">Evaluation Trace Log:</div>
                    {% for line in results.debug_trace %}
                    <div class="py-1 border-b border-gray-200 last:border-0">{{ line }}</div>
                    {% endfor %}
                </div>
            </div>
        </div>
        {% endif %}
    </div>
<!--
    <footer class="bg-white border-t border-gray-200 py-6 mt-12">
        <div class="container mx-auto px-4 text-center text-gray-500 text-sm">
            <p>Academic Project Report Evaluator | Powered by AI Technology</p>
          
        </div>
    </footer>
-->
    <script>
        // File upload display
        const fileInput = document.getElementById('fileInput');
        const fileNameDisplay = document.getElementById('fileNameDisplay');
        const fileUpload = document.querySelector('.file-upload');
        
        fileInput.addEventListener('change', (e) => {
            if (e.target.files.length) {
                fileNameDisplay.textContent = `Selected: ${e.target.files[0].name}`;

                fileNameDisplay.classList.remove('hidden');
            }
        });
        
        // Drag and drop
        ['dragenter', 'dragover', 'dragleave', 'drop'].forEach(eventName => {
            fileUpload.addEventListener(eventName, preventDefaults, false);
        });
        
        function preventDefaults(e) {
            e.preventDefault();
            e.stopPropagation();
        }
        
        ['dragenter', 'dragover'].forEach(eventName => {
            fileUpload.addEventListener(eventName, highlight, false);
        });
        
        ['dragleave', 'drop'].forEach(eventName => {
            fileUpload.addEventListener(eventName, unhighlight, false);
        });
        
        function highlight() {
            fileUpload.classList.add('dragover');
        }
        
        function unhighlight() {
            fileUpload.classList.remove('dragover');
        }
        
        fileUpload.addEventListener('drop', handleDrop, false);
        
        function handleDrop(e) {
            const dt = e.dataTransfer;
            const files = dt.files;
            fileInput.files = files;
            fileNameDisplay.textContent = `Selected: ${files[0].name}`;

            fileNameDisplay.classList.remove('hidden');
        }
        
        // Form submission spinner
        const form = document.getElementById('evaluationForm');
        const spinner = document.getElementById('spinner');
        const submitText = document.getElementById('submitText');
        
        form.addEventListener('submit', () => {
            spinner.style.display = 'inline-block';
            submitText.textContent = 'Processing...';
        });
        
        // Debug toggle
        function toggleDebug() {
            const debugInfo = document.getElementById('debugInfo');
            const debugToggleText = document.getElementById('debugToggleText');
            debugInfo.classList.toggle('hidden');
            debugToggleText.textContent = debugInfo.classList.contains('hidden') ? 
                'Show Debug Information' : 'Hide Debug Information';
        }
        
        // Add current datetime to the page
        document.addEventListener('DOMContentLoaded', function() {
            const now = new Date();
            const datetimeElement = document.createElement('div');
            datetimeElement.className = 'hidden';
            datetimeElement.id = 'currentDatetime';
            datetimeElement.textContent = now.toISOString();
            document.body.appendChild(datetimeElement);
        });
    </script>
</body>
</html>
"""

# ========== Main Execution ==========
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
